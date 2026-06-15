"""Generate IEEE-format literature review on Agentic AI."""

from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Cm

PROJECT_DIR = Path(__file__).parent
CHARTS_DIR = PROJECT_DIR / "charts"
OUTPUT = PROJECT_DIR / "ai.docx"

IEEE_FONT = "Times New Roman"
BODY_SIZE = Pt(10)
TITLE_SIZE = Pt(24)
SECTION_SIZE = Pt(10)
AUTHOR_SIZE = Pt(11)


def set_two_columns(section):
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "360")


def fmt_run(run, size=BODY_SIZE, bold=False, italic=False, font=IEEE_FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic


def add_para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, italic=False, size=BODY_SIZE, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    fmt_run(run, size=size, bold=bold, italic=italic)
    return p


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    fmt_run(run, size=SECTION_SIZE, bold=True)
    return p


def add_subsection_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    fmt_run(run, size=BODY_SIZE, bold=True, italic=True)
    return p


def add_ieee_table(doc, caption, headers, rows):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    fmt_run(cr, size=Pt(9), italic=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                fmt_run(r, bold=True, size=Pt(9))
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for r in p.runs:
                    fmt_run(r, size=Pt(9))
    doc.add_paragraph()
    return table


def chart_agentic_taxonomy():
    fig, ax = plt.subplots(figsize=(7.5, 4.8))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7)
    ax.axis("off")

    boxes = {
        "Agentic AI": (3.5, 5.8, 3, 0.7, "#1E3A5F"),
        "Reasoning": (0.3, 3.8, 2.2, 0.65, "#2563EB"),
        "Acting": (3.9, 3.8, 2.2, 0.65, "#7C3AED"),
        "Interacting": (7.5, 3.8, 2.2, 0.65, "#0891B2"),
        "Planning &\nReflection": (0.1, 2.0, 2.6, 0.8, "#93C5FD"),
        "Memory &\nRetrieval": (0.1, 0.8, 2.6, 0.8, "#93C5FD"),
        "Tool Use &\nMCP": (3.7, 2.0, 2.6, 0.8, "#C4B5FD"),
        "Orchestration": (3.7, 0.8, 2.6, 0.8, "#C4B5FD"),
        "Multi-Agent\nCollaboration": (7.3, 2.0, 2.6, 0.8, "#67E8F9"),
        "Human-in-\nthe-Loop": (7.3, 0.8, 2.6, 0.8, "#67E8F9"),
    }

    for label, (x, y, w, h, color) in boxes.items():
        rect = mpatches.FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.02",
            facecolor=color, edgecolor="white", linewidth=1.5, alpha=0.92
        )
        ax.add_patch(rect)
        weight = "bold" if label == "Agentic AI" else "normal"
        fsize = 11 if label == "Agentic AI" else 8.5
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center",
                fontsize=fsize, fontweight=weight, color="white" if label == "Agentic AI" else "#1E293B")

    arrows = [
        ((5, 5.8), (1.4, 4.5)), ((5, 5.8), (5, 4.5)), ((5, 5.8), (8.6, 4.5)),
        ((1.4, 3.8), (1.4, 2.8)), ((1.4, 3.8), (1.4, 1.6)),
        ((5, 3.8), (5, 2.8)), ((5, 3.8), (5, 1.6)),
        ((8.6, 3.8), (8.6, 2.8)), ((8.6, 3.8), (8.6, 1.6)),
    ]
    for (x1, y1), (x2, y2) in arrows:
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="->", color="#64748B", lw=1.2))

    ax.set_title("Fig. 1. Taxonomy of Agentic AI Capabilities", fontsize=12, fontweight="bold", pad=10)
    plt.tight_layout()
    path = CHARTS_DIR / "agentic_taxonomy.png"
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def chart_timeline():
    events = [
        ("ReAct", "2022"),
        ("Toolformer", "2023"),
        ("AutoGPT", "2023"),
        ("AgentBench", "2023"),
        ("LangGraph", "2024"),
        ("MCP", "2024"),
        ("GAIA", "2024"),
        ("Multi-Agent\nSurveys", "2025"),
    ]
    fig, ax = plt.subplots(figsize=(7.5, 2.8))
    years = [int(e[1]) for e in events]
    labels = [e[0] for e in events]
    y_pos = [0] * len(events)
    ax.hlines(0, 2021.5, 2025.8, colors="#CBD5E1", linewidth=3)
    ax.scatter(years, y_pos, s=120, c="#2563EB", zorder=3, edgecolors="white", linewidth=1.5)
    for i, (label, year) in enumerate(events):
        offset = 0.35 if i % 2 == 0 else -0.55
        ax.annotate(f"{label}\n({year})", (int(year), 0), textcoords="offset points",
                    xytext=(0, 30 if offset > 0 else -40), ha="center", fontsize=8, fontweight="bold")
    ax.set_yticks([])
    ax.set_xlim(2021.5, 2025.8)
    ax.set_xlabel("Year", fontsize=10)
    ax.set_title("Fig. 2. Milestones in Agentic AI Research and Frameworks", fontsize=11, fontweight="bold")
    ax.spines["left"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["top"].set_visible(False)
    plt.tight_layout()
    path = CHARTS_DIR / "agentic_timeline.png"
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def chart_framework_comparison():
    frameworks = ["LangGraph", "AutoGen", "CrewAI", "MetaGPT", "CAMEL", "smolagents"]
    scores = [9, 8, 7, 8, 7, 6]  # illustrative adoption/flexibility index
    fig, ax = plt.subplots(figsize=(7.5, 3.5))
    colors = ["#2563EB", "#7C3AED", "#0891B2", "#059669", "#D97706", "#DC2626"]
    bars = ax.barh(frameworks, scores, color=colors, edgecolor="white", height=0.6)
    for bar, s in zip(bars, scores):
        ax.text(bar.get_width() + 0.1, bar.get_y() + bar.get_height() / 2,
                f"{s}/10", va="center", fontsize=9, fontweight="bold")
    ax.set_xlim(0, 11)
    ax.set_xlabel("Ecosystem Maturity Index (illustrative)", fontsize=10)
    ax.set_title("Fig. 3. Comparative Maturity of Agent Frameworks", fontsize=11, fontweight="bold")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.invert_yaxis()
    plt.tight_layout()
    path = CHARTS_DIR / "framework_comparison.png"
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def chart_agent_loop():
    fig, ax = plt.subplots(figsize=(5, 5))
    ax.set_xlim(-1.5, 1.5)
    ax.set_ylim(-1.5, 1.5)
    ax.set_aspect("equal")
    ax.axis("off")
    steps = ["Perceive", "Reason", "Plan", "Act", "Observe", "Reflect"]
    n = len(steps)
    import numpy as np
    angles = np.linspace(90, 90 - 360, n, endpoint=False) * np.pi / 180
    r = 1.0
    xs = r * np.cos(angles)
    ys = r * np.sin(angles)
    for i in range(n):
        j = (i + 1) % n
        ax.annotate("", xy=(xs[j], ys[j]), xytext=(xs[i], ys[i]),
                    arrowprops=dict(arrowstyle="->", color="#64748B", lw=1.8,
                                    connectionstyle="arc3,rad=0.15"))
    for i, (x, y, label) in enumerate(zip(xs, ys, steps)):
        circle = plt.Circle((x, y), 0.32, color="#2563EB", ec="white", lw=2, zorder=3)
        ax.add_patch(circle)
        ax.text(x, y, label, ha="center", va="center", fontsize=7.5,
                fontweight="bold", color="white", zorder=4)
    ax.text(0, 0, "Agent\nLoop", ha="center", va="center", fontsize=11,
            fontweight="bold", color="#1E293B")
    ax.set_title("Fig. 4. Canonical Agent Control Loop", fontsize=11, fontweight="bold", pad=12)
    plt.tight_layout()
    path = CHARTS_DIR / "agent_loop.png"
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def build_paper():
    CHARTS_DIR.mkdir(exist_ok=True)
    tax_chart = chart_agentic_taxonomy()
    timeline_chart = chart_timeline()
    framework_chart = chart_framework_comparison()
    loop_chart = chart_agent_loop()

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.left_margin = Cm(1.78)
    section.right_margin = Cm(1.78)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.54)
    set_two_columns(section)

    # Title
    add_para(doc, "Agentic Artificial Intelligence: A Literature Review",
             align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=TITLE_SIZE, space_after=8)

    add_para(doc, "Research Student",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=AUTHOR_SIZE, italic=True, space_after=2)
    add_para(doc, "Department of Computer Science and Engineering",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(9), italic=True, space_after=2)
    add_para(doc, "Email: student@university.edu",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(9), italic=True, space_after=12)

    # Abstract
    abs_p = doc.add_paragraph()
    abs_p.paragraph_format.space_after = Pt(6)
    abs_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = abs_p.add_run("Abstract—")
    fmt_run(r1, bold=True, italic=True)
    r2 = abs_p.add_run(
        "Agentic artificial intelligence (AI) refers to systems that extend large language models "
        "(LLMs) with autonomous planning, tool use, memory, and multi-step interaction to pursue "
        "complex goals in dynamic environments. This literature review synthesizes recent academic "
        "and industry research on agentic AI, spanning foundational paradigms such as ReAct and "
        "Toolformer, architectural frameworks including LangGraph and AutoGen, evaluation benchmarks "
        "such as AgentBench and GAIA, and emerging protocols for tool connectivity and multi-agent "
        "coordination. We organize the literature along three axes—reasoning, acting, and "
        "interacting—and discuss open challenges in evaluation methodology, safety, cost efficiency, "
        "and reproducibility. The review concludes with research directions for building reliable, "
        "controllable, and production-ready agentic systems."
    )
    fmt_run(r2)

    # Index Terms
    kw_p = doc.add_paragraph()
    kw_p.paragraph_format.space_after = Pt(10)
    kw_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    kr = kw_p.add_run("Index Terms—")
    fmt_run(kr, bold=True, italic=True)
    kv = kw_p.add_run(
        "Agentic AI, large language models, autonomous agents, ReAct, multi-agent systems, "
        "tool use, evaluation benchmarks, literature review."
    )
    fmt_run(kv)

    # I. INTRODUCTION
    add_section_heading(doc, "I. INTRODUCTION")
    add_para(doc,
        "The emergence of large language models (LLMs) has catalyzed a paradigm shift from "
        "passive text generation toward autonomous, goal-directed artificial intelligence. "
        "Agentic AI denotes systems in which an LLM serves as a cognitive controller that "
        "perceives context, reasons about goals, selects actions, invokes external tools, "
        "and adapts behavior through environmental feedback over extended trajectories [1], [2]. "
        "Unlike traditional NLP pipelines that map inputs to single outputs, agentic systems "
        "operate in closed-loop control cycles, pursuing multi-step objectives in software "
        "environments, web interfaces, scientific workflows, and enterprise automation [3]."
    )
    add_para(doc,
        "This review addresses the rapid proliferation of surveys, benchmarks, and frameworks "
        "that has nonetheless left the field fragmented. Wang et al. [1] provide a holistic "
        "survey of LLM-based autonomous agents, while Pati [4] offers a comprehensive treatment "
        "of technologies and societal implications published in IEEE Access. More recent work "
        "by Xi et al. [2] categorizes agentic LLMs along reasoning, acting, and interacting "
        "dimensions, and Guo et al. [5] formalize agentic architectures using partially "
        "observable Markov decision process (POMDP) control loops. Our contribution is to "
        "synthesize these perspectives into a unified literature review suitable for researchers "
        "entering the field, with emphasis on architectural patterns, evaluation practice, and "
        "deployment challenges. An interactive companion guide (agentic_ai_explorer.html) provides "
        "curated papers, frameworks, and a structured learning path."
    )

    # II. RELATED WORK
    add_section_heading(doc, "II. RELATED WORK")
    add_para(doc,
        "Prior surveys address complementary facets of the agentic AI landscape. Wang et al. [1] "
        "survey LLM-based autonomous agents with emphasis on construction and applications. Xi et al. "
        "[2] organize agentic LLMs along reasoning, acting, and interacting axes. Pati [4] publishes "
        "a broad IEEE Access survey covering societal implications. Guo et al. [5] adopt an "
        "engineering-focused POMDP control-loop perspective. Chen et al. [20] review multi-agent "
        "collaboration, while recent evaluation surveys [3], [7] stress multi-dimensional benchmarking. "
        "Table III positions this review relative to prior work."
    )
    add_ieee_table(doc, "TABLE III\nCOMPARISON WITH PRIOR SURVEYS",
        ["Survey", "Focus", "Year"],
        [
            ["Wang et al. [1]", "Construction & applications", "2024"],
            ["Xi et al. [2]", "Reasoning/acting/interacting", "2025"],
            ["Pati [4]", "Technologies & society (IEEE)", "2025"],
            ["Guo et al. [5]", "Architecture & POMDP loop", "2026"],
            ["This review", "Unified synthesis + evaluation", "2026"],
        ]
    )

    # III. BACKGROUND AND DEFINITIONS
    add_section_heading(doc, "III. BACKGROUND AND DEFINITIONS")
    add_subsection_heading(doc, "A. From LLMs to Agents")
    add_para(doc,
        "Classical autonomous agents in AI research operated within isolated environments with "
        "hand-crafted knowledge representations [1]. LLMs pretrained on web-scale corpora provide "
        "a qualitatively different foundation: broad world knowledge, in-context learning, and "
        "natural-language interfaces that lower the barrier to tool integration [6]. An LLM agent "
        "typically comprises four modules: a profile (role and constraints), memory (short- and "
        "long-term), planning (decomposition and reflection), and action (tool invocation and "
        "environment interaction) [1], [5]."
    )
    add_subsection_heading(doc, "B. Agentic AI vs. Generative AI")
    add_para(doc,
        "Generative AI focuses on producing content—text, code, or images—from prompts. Agentic "
        "AI extends generation with agency: the capacity to observe state changes, revise plans, "
        "and execute sequences of operations until a termination condition is met [4], [7]. "
        "Recent surveys argue that enterprise adoption hinges on evaluating not only output "
        "quality but also latency, cost, reliability, safety, and multi-agent coordination under "
        "production constraints [7], [8]."
    )

    add_subsection_heading(doc, "C. Agent Environments")
    add_para(doc,
        "Agents operate in environments classified along multiple axes: static versus dynamic, "
        "fully versus partially observable, deterministic versus stochastic, and single- versus "
        "multi-agent [30]. Environmental complexity directly affects planning horizon, tool "
        "selection, and the need for memory and reflection. Web-browsing agents face stochastic, "
        "partially observable environments; code agents operate in deterministic sandboxes with "
        "clear termination conditions [22], [24]."
    )

    # IV. FOUNDATIONAL PARADIGMS
    add_section_heading(doc, "IV. FOUNDATIONAL PARADIGMS")
    add_subsection_heading(doc, "A. ReAct: Reasoning and Acting")
    add_para(doc,
        "Yao et al. [9] introduced ReAct (Reasoning + Acting), interleaving chain-of-thought "
        "rationales with executable actions and environment observations. The trajectory "
        "τ = {o₁, z₁, a₁, o₂, z₂, a₂, …} grounds reasoning in external feedback, reducing "
        "hallucination compared to latent chain-of-thought prompting alone [5], [9]. ReAct "
        "remains the de facto scaffold for modern agent benchmarks and frameworks, though it is "
        "vulnerable to error propagation and myopic planning when early mistakes trigger "
        "unproductive loops [5]."
    )
    add_subsection_heading(doc, "B. Tool Use and Function Calling")
    add_para(doc,
        "Schick et al. [10] proposed Toolformer, demonstrating that LMs can learn to invoke APIs "
        "through self-supervised fine-tuning. Subsequent systems expose tools via JSON schemas, "
        "OpenAI function calling, and standardized protocols such as the Model Context Protocol "
        "(MCP), which decouples tool providers from agent runtimes [11]. Tool-augmented agents "
        "enable retrieval-augmented generation (RAG), code execution, database queries, and web "
        "browsing, substantially expanding the feasible task space [2], [12]."
    )
    add_subsection_heading(doc, "C. Reflection and Self-Improvement")
    add_para(doc,
        "Reflexion [13] extends ReAct with an evaluator–reflector loop that generates verbal "
        "reinforcement cues from failed trajectories, enabling iterative self-correction without "
        "weight updates. Self-Refine and related methods apply similar actor–critic decompositions "
        "at inference time [2]. These approaches improve robustness on coding and reasoning tasks "
        "but increase token cost and latency, raising questions about cost-efficiency in "
        "production deployments [14]."
    )

    add_ieee_table(doc, "TABLE I\nCOMPARISON OF FOUNDATIONAL AGENTIC PARADIGMS",
        ["Paradigm", "Year", "Key Idea", "Limitation"],
        [
            ["ReAct [9]", "2022", "Interleave thought and action", "Error propagation"],
            ["Toolformer [10]", "2023", "Self-supervised tool learning", "Training overhead"],
            ["Reflexion [13]", "2023", "Verbal self-reflection", "Higher inference cost"],
            ["AutoGPT [15]", "2023", "Autonomous goal loops", "Stability concerns"],
        ]
    )

    add_subsection_heading(doc, "D. Memory Architectures")
    add_para(doc,
        "Long-horizon agent tasks require memory beyond the context window. Hierarchical memory "
        "systems combine working memory (current trajectory), episodic memory (past interactions), "
        "and semantic memory (retrieved knowledge) [2], [31]. Frameworks such as Mem0, Letta, and "
        "LangChain memory modules implement vector-store retrieval, graph-based knowledge, and "
        "self-managed memory augmentation [31]. MemInsight demonstrates that autonomous memory "
        "tagging and filtering can substantially improve multi-session task performance."
    )

    # V. ARCHITECTURES AND FRAMEWORKS
    add_section_heading(doc, "V. ARCHITECTURES AND FRAMEWORKS")
    add_subsection_heading(doc, "A. Single-Agent Orchestration")
    add_para(doc,
        "Community frameworks such as AutoGPT and BabyAGI popularized fully autonomous loops "
        "with goal decomposition and persistent memory [1], [15]. LangChain and LangGraph "
        "introduced controllable orchestration through explicit state graphs, enabling developers "
        "to define nodes, edges, and conditional transitions rather than unconstrained LLM loops "
        "[5], [11]. Guo et al. [5] contrast autonomous loops with controllable graphs, arguing "
        "that production systems require explicit workflow structure to manage failure modes."
    )
    add_subsection_heading(doc, "B. Multi-Agent Systems")
    add_para(doc,
        "Multi-agent LLM systems assign specialized roles—planner, coder, critic—to collaborating "
        "agents. Frameworks including CAMEL [16], AutoGen [17], MetaGPT [18], and CrewAI implement "
        "topologies such as chain, star, mesh, and hierarchical workflows [5], [19]. Chen et al. [20] "
        "survey multi-agent collaboration mechanisms, while MASEval [21] demonstrates that framework "
        "choice can impact performance as much as backbone model selection within a capability tier."
    )
    add_subsection_heading(doc, "C. Agent Protocols")
    add_para(doc,
        "Interoperability standards are emerging to connect agents with tools and other agents. "
        "The Model Context Protocol (MCP) standardizes tool and resource discovery; the Agent "
        "Communication Protocol (ACP) and Agent-to-Agent Protocol (A2A) address cross-agent "
        "messaging [11], [12]. These protocols reduce integration friction but introduce new "
        "attack surfaces related to prompt injection and unauthorized tool invocation [12]."
    )

  # Figure 1 - need to span columns - in two column mode, add picture in single column
    fig_p = doc.add_paragraph()
    fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig_p.add_run()
    run.add_picture(str(tax_chart), width=Inches(3.2))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run("Fig. 1. Taxonomy of agentic AI organized along reasoning, acting, and interacting axes.")
    fmt_run(cr, size=Pt(9), italic=True)

    add_ieee_table(doc, "TABLE IV\nAGENT FRAMEWORK COMPARISON",
        ["Framework", "Paradigm", "Multi-Agent", "Production"],
        [
            ["LangGraph", "State graph", "Yes", "High"],
            ["AutoGen", "Conversation", "Yes", "Medium"],
            ["CrewAI", "Role-based teams", "Yes", "Medium"],
            ["MetaGPT", "SOP simulation", "Yes", "Medium"],
            ["smolagents", "Code-as-action", "Limited", "Emerging"],
        ]
    )

    fig_loop = doc.add_paragraph()
    fig_loop.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_loop.add_run().add_picture(str(loop_chart), width=Inches(2.2))
    cap_loop = doc.add_paragraph()
    cap_loop.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr_loop = cap_loop.add_run("Fig. 2. Canonical perceive–reason–plan–act–observe–reflect control loop.")
    fmt_run(cr_loop, size=Pt(9), italic=True)

    # VI. EVALUATION AND BENCHMARKS
    add_section_heading(doc, "VI. EVALUATION AND BENCHMARKS")
    add_subsection_heading(doc, "A. Benchmark Landscape")
    add_para(doc,
        "Evaluation has shifted from static NLP metrics toward interactive, environment-grounded "
        "assessments. Liu et al. [22] introduced AgentBench spanning eight environments including "
        "operating systems, databases, and web shopping. Mialon et al. [23] proposed GAIA for "
        "general AI assistants requiring tool use and multi-step reasoning. Jimenez et al. [24] "
        "released SWE-bench for software engineering agents, while Yao et al. [25] developed "
        "τ-bench for conversational service tasks with simulated users [3], [7]."
    )
    add_subsection_heading(doc, "B. Evaluation Gaps")
    add_para(doc,
        "Despite benchmark proliferation, several gaps persist. First, most evaluations are "
        "model-centric, conflating backbone LLM capability with framework implementation [21], "
        "[14]. Second, laboratory benchmarks underrepresent cost, latency, and reliability "
        "requirements of enterprise deployment [7], [8]. Third, multi-dimensional evaluation "
        "taxonomies—covering safety, factuality, coordination, and tool correctness—remain "
        "inconsistently applied [3], [26]. Ling et al. [3] propose a unified framework "
        "standardizing seven benchmarks as instruction–tool–environment triplets under a fixed "
        "ReAct scaffold to isolate capability measurement."
    )

    add_ieee_table(doc, "TABLE II\nREPRESENTATIVE AGENTIC AI BENCHMARKS",
        ["Benchmark", "Domain", "Key Metric", "Ref."],
        [
            ["AgentBench", "Multi-environment", "Success rate", "[22]"],
            ["GAIA", "General assistance", "Level-wise accuracy", "[23]"],
            ["SWE-bench", "Software engineering", "Issue resolution", "[24]"],
            ["τ-bench", "Customer service", "Task completion", "[25]"],
            ["BFCL", "Function calling", "Tool invocation accuracy", "[27]"],
            ["MASEval", "Multi-agent systems", "System-level score", "[21]"],
        ]
    )

    # VII. APPLICATIONS
    add_section_heading(doc, "VII. APPLICATIONS")
    add_para(doc,
        "Agentic AI has been applied across scientific discovery, software engineering, healthcare, "
        "finance, and education [1], [4], [12]. Wang et al. [1] catalog applications in social "
        "science (simulation and survey automation), natural science (experiment planning and "
        "literature synthesis), and engineering (code generation and DevOps). In software "
        "engineering, agents such as SWE-Agent [28] demonstrate that specialized interfaces "
        "and ReAct-style loops can resolve real GitHub issues. Biomedical and materials-science "
        "agents integrate domain tools for hypothesis generation and structured data retrieval "
        "[12]. Nevertheless, domain-specific customization often outperforms general-purpose "
        "agents on specialized benchmarks, though recent studies show narrowing gaps on several "
        "tasks [14]."
    )

    # VIII. CHALLENGES
    add_section_heading(doc, "VIII. CHALLENGES AND OPEN PROBLEMS")
    add_subsection_heading(doc, "A. Reliability and Safety")
    add_para(doc,
        "Agentic systems amplify LLM failure modes: hallucinated tool calls, infinite loops, "
        "and cascading errors across multi-step trajectories [5], [9]. Safety benchmarks including "
        "Agent-SafetyBench [29] and red-teaming studies reveal vulnerabilities to prompt "
        "injection, data exfiltration via tools, and unsafe action sequences [12]. Production "
        "systems require guardrails, human-in-the-loop approval, and formal verification of "
        "critical actions [4], [7]."
    )
    add_subsection_heading(doc, "B. Cost and Efficiency")
    add_para(doc,
        "Each reasoning step, reflection cycle, and tool call incurs API latency and token cost. "
        "Exgentic [14] reports wide variation in cost-per-task across agent architectures and "
        "notes that open-weight models exhibit \"generality sinks\"—collapse on specific "
        "architectures or benchmarks absent in frontier closed-source models. Efficient agent "
        "design must jointly optimize accuracy, latency, and dollar cost [7], [8]."
    )
    add_subsection_heading(doc, "C. Reproducibility")
    add_para(doc,
        "Non-deterministic sampling, evolving APIs, and live web environments hinder reproducible "
        "evaluation [3], [7]. Recent frameworks adopt static environment snapshots, fixed "
        "scaffolds, and open trace logging to improve comparability [3], [14], [21]."
    )

    # Timeline figure
    fig_p2 = doc.add_paragraph()
    fig_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = fig_p2.add_run()
    run2.add_picture(str(timeline_chart), width=Inches(3.2))
    cap2 = doc.add_paragraph()
    cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr2 = cap2.add_run("Fig. 3. Selected milestones in agentic AI research (2022–2025).")
    fmt_run(cr2, size=Pt(9), italic=True)

    add_subsection_heading(doc, "D. Societal and Ethical Implications")
    add_para(doc,
        "Pati [4] and related work highlight alignment risks, accountability gaps, and labor "
        "displacement concerns as agents gain autonomy in customer service, software engineering, "
        "and scientific research. Dual-use capabilities—autonomous code execution, web access, and "
        "data exfiltration via tools—require governance frameworks, audit trails, and human oversight "
        "for high-stakes decisions [4], [29]."
    )

    fig_p3 = doc.add_paragraph()
    fig_p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_p3.add_run().add_picture(str(framework_chart), width=Inches(3.2))
    cap3 = doc.add_paragraph()
    cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr3 = cap3.add_run("Fig. 4. Illustrative framework ecosystem maturity comparison.")
    fmt_run(cr3, size=Pt(9), italic=True)

    # IX. FUTURE DIRECTIONS
    add_section_heading(doc, "IX. FUTURE RESEARCH DIRECTIONS")
    add_para(doc,
        "The literature points to several high-impact research directions: (1) system-level "
        "evaluation that treats framework topology and orchestration as first-class variables "
        "[21]; (2) unified, multi-dimensional benchmarking spanning safety, cost, and "
        "coordination [7], [8]; (3) reinforcement learning for dynamic tool selection and "
        "long-horizon planning [12]; (4) standardized agent protocols with formal security "
        "models [11]; (5) neuro-symbolic integration for verifiable reasoning in high-stakes "
        "domains [4]; and (6) human-agent collaboration frameworks that balance autonomy with "
        "accountability [2], [4]."
    )

    # X. CONCLUSION
    add_section_heading(doc, "X. CONCLUSION")
    add_para(doc,
        "Agentic AI represents a maturing research field that unites LLM reasoning, tool-augmented "
        "action, and multi-agent interaction into deployable systems. Foundational paradigms from "
        "ReAct to Reflexion established the inference-time control loop; frameworks from LangGraph "
        "to AutoGen operationalized multi-agent workflows; and benchmarks from AgentBench to GAIA "
        "began measuring real-world competence. Yet the transition from impressive demonstrations "
        "to dependable production systems requires advances in system-level evaluation, safety "
        "engineering, and cost-aware architecture design. This review provides a structured entry "
        "point for researchers and practitioners navigating this rapidly evolving landscape."
    )

    add_section_heading(doc, "ACKNOWLEDGMENT")
    add_para(doc,
        "The author thanks the open research community for maintaining survey repositories and "
        "benchmark suites that made this review possible. An interactive learning companion "
        "(agentic_ai_explorer.html) accompanies this paper."
    )

    # REFERENCES
    add_section_heading(doc, "REFERENCES")
    references = [
        '[1] L. Wang et al., "A survey on large language model based autonomous agents," Frontiers of Computer Science, vol. 18, no. 6, Art. no. 186345, Dec. 2024.',
        '[2] Y. Xi et al., "Agentic large language models, a survey," arXiv preprint arXiv:2503.23037, 2025.',
        '[3] Y. Ling et al., "A unified framework for the evaluation of LLM agentic capabilities," arXiv preprint arXiv:2605.27898, 2026.',
        '[4] A. K. Pati, "Agentic AI: A comprehensive survey of technologies, applications, and societal implications," IEEE Access, vol. 13, pp. 151824–151837, 2025.',
        '[5] Y. Guo et al., "Agentic artificial intelligence (AI): Architectures, taxonomies, and evaluation of large language model agents," arXiv preprint arXiv:2601.12560, 2026.',
        '[6] T. Brown et al., "Language models are few-shot learners," in Proc. Adv. Neural Inf. Process. Syst. (NeurIPS), 2020.',
        '[7] Anonymous, "Evaluation and benchmarking of generative and agentic AI systems: A comprehensive survey," Preprints.org, 2025.',
        '[8] Stanford HAI, "Artificial Intelligence Index Report 2025," Stanford Univ., Stanford, CA, USA, 2025.',
        '[9] S. Yao et al., "ReAct: Synergizing reasoning and acting in language models," arXiv preprint arXiv:2210.03629, 2022.',
        '[10] T. Schick et al., "Toolformer: Language models can teach themselves to use tools," in Proc. Adv. Neural Inf. Process. Syst. (NeurIPS), 2023.',
        '[11] Anthropic, "Model Context Protocol specification," 2024. [Online]. Available: https://modelcontextprotocol.io',
        '[12] H. Wang et al., "From LLM reasoning to autonomous AI agents: A comprehensive review," arXiv preprint arXiv:2504.19678, 2025.',
        '[13] N. Shinn et al., "Reflexion: Language agents with verbal reinforcement learning," in Proc. Adv. Neural Inf. Process. Syst. (NeurIPS), 2023.',
        '[14] Exgentic Research, "General agent evaluation," arXiv preprint arXiv:2602.22953, 2026.',
        '[15] Significant Gravitas, "AutoGPT: An autonomous GPT-4 experiment," GitHub repository, 2023.',
        '[16] G. Li et al., "CAMEL: Communicative agents for mind exploration of large language model society," in Proc. Adv. Neural Inf. Process. Syst. (NeurIPS), 2023.',
        '[17] Q. Wu et al., "AutoGen: Enabling next-gen LLM applications via multi-agent conversation," arXiv preprint arXiv:2308.08155, 2023.',
        '[18] S. Hong et al., "MetaGPT: Meta programming for multi-agent collaborative framework," in Proc. Int. Conf. Learn. Represent. (ICLR), 2024.',
        '[19] J. Xi et al., "The rise and potential of large language model based agents: A survey," arXiv preprint arXiv:2309.07864, 2023.',
        '[20] L. Chen et al., "Large language model-based multi-agents: A survey of progress and challenges," in Proc. IJCAI, 2024, pp. 8048–8057.',
        '[21] O. Wiest et al., "MASEval: Extending multi-agent evaluation from models to systems," arXiv preprint arXiv:2603.08835, 2026.',
        '[22] X. Liu et al., "AgentBench: Evaluating LLMs as agents," in Proc. Int. Conf. Learn. Represent. (ICLR), 2024.',
        '[23] G. Mialon et al., "GAIA: A benchmark for general AI assistants," in Proc. Int. Conf. Learn. Represent. (ICLR), 2024.',
        '[24] C. Jimenez et al., "SWE-bench: Can language models resolve real-world GitHub issues?" in Proc. Int. Conf. Learn. Represent. (ICLR), 2024.',
        '[25] S. Yao et al., "τ-bench: A benchmark for tool-agent-user interaction in real-world domains," arXiv preprint arXiv:2410.24185, 2024.',
        '[26] Y. Xi et al., "The rise and potential of large language model based agents: A survey," ACM Comput. Surv., 2025.',
        '[27] S. G. Patil et al., "Gorilla: Large language model connected with massive APIs," in Proc. Adv. Neural Inf. Process. Syst. (NeurIPS), 2024.',
        '[28] J. Yang et al., "SWE-agent: Agent-computer interfaces enable automated software engineering," arXiv preprint arXiv:2405.15793, 2024.',
        '[29] Z. Zhang et al., "Agent-safetybench: Evaluating the safety of LLM agents," arXiv preprint arXiv:2412.13178, 2024.',
        '[30] S. Russell and P. Norvig, Artificial Intelligence: A Modern Approach, 4th ed. Pearson, 2020.',
        '[31] Y. Wang et al., "MemInsight: Autonomous memory augmentation for LLM agents," arXiv preprint arXiv:2503.21760, 2025.',
    ]

    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(14)
        p.paragraph_format.first_line_indent = Pt(-14)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(ref)
        fmt_run(run, size=Pt(9))

    doc.save(OUTPUT)
    print(f"Paper saved: {OUTPUT}")
    print(f"Size: {OUTPUT.stat().st_size:,} bytes")


if __name__ == "__main__":
    build_paper()
