"""Generate SpaceX stock information report with tables and figures."""

from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_DIR = Path(__file__).parent
CHARTS_DIR = PROJECT_DIR / "charts"
OUTPUT = PROJECT_DIR / "spacex.docx"


def set_cell_shading(cell, color_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, header_color="1F4E79"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], header_color)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    for row_data in rows:
        row = table.add_row().cells
        for i, text in enumerate(row_data):
            row[i].text = str(text)
    return table


def chart_stock_price():
  dates = ["Jun 11\n(Private)", "Jun 12\n(IPO Open)", "Jun 12\n(Close)", "Jun 15\n(Current)"]
  prices = [141.09, 150.00, 160.95, 171.91]
  fig, ax = plt.subplots(figsize=(8, 4.5))
  bars = ax.bar(dates, prices, color=["#6B7280", "#3B82F6", "#2563EB", "#1D4ED8"], edgecolor="white", linewidth=1.2)
  ax.axhline(y=135, color="#EF4444", linestyle="--", linewidth=1.5, label="IPO Price ($135)")
  for bar, price in zip(bars, prices):
    ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 2, f"${price:.2f}",
            ha="center", va="bottom", fontsize=10, fontweight="bold")
  ax.set_ylabel("Share Price (USD)", fontsize=11)
  ax.set_title("SpaceX (SPCX) Stock Price — IPO Week (June 2026)", fontsize=13, fontweight="bold", pad=12)
  ax.yaxis.set_major_formatter(mticker.FormatStrFormatter("$%.0f"))
  ax.set_ylim(0, 200)
  ax.legend(loc="upper left")
  ax.spines["top"].set_visible(False)
  ax.spines["right"].set_visible(False)
  ax.grid(axis="y", alpha=0.3)
  plt.tight_layout()
  path = CHARTS_DIR / "stock_price.png"
  fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
  plt.close(fig)
  return path


def chart_revenue_breakdown():
  labels = ["Starlink", "Launch Services", "Other / xAI"]
  values = [11.4, 4.8, 2.5]
  colors = ["#0EA5E9", "#8B5CF6", "#F59E0B"]
  fig, ax = plt.subplots(figsize=(7, 5))
  wedges, texts, autotexts = ax.pie(
    values, labels=labels, autopct="%1.1f%%", startangle=140,
    colors=colors, explode=(0.04, 0, 0), pctdistance=0.78,
    textprops={"fontsize": 10}
  )
  for t in autotexts:
    t.set_fontweight("bold")
  centre = plt.Circle((0, 0), 0.55, fc="white")
  ax.add_artist(centre)
  ax.text(0, 0, "$18.7B\nTotal", ha="center", va="center", fontsize=14, fontweight="bold")
  ax.set_title("SpaceX 2025 Revenue by Segment", fontsize=13, fontweight="bold", pad=16)
  plt.tight_layout()
  path = CHARTS_DIR / "revenue_breakdown.png"
  fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
  plt.close(fig)
  return path


def chart_valuation_timeline():
  periods = ["Jul 2025", "Dec 2025", "Feb 2026\n(xAI merger)", "Jun 2026\n(IPO)", "Jun 15\n(Market)"]
  valuations = [400, 800, 1250, 1750, 2100]
  fig, ax = plt.subplots(figsize=(8, 4.5))
  ax.plot(periods, valuations, marker="o", linewidth=2.5, color="#1E40AF", markersize=9, markerfacecolor="#3B82F6")
  ax.fill_between(range(len(periods)), valuations, alpha=0.12, color="#3B82F6")
  for i, v in enumerate(valuations):
    ax.annotate(f"${v}B", (i, v), textcoords="offset points", xytext=(0, 10),
                ha="center", fontsize=9, fontweight="bold")
  ax.set_ylabel("Valuation (USD Billions)", fontsize=11)
  ax.set_title("SpaceX Valuation Timeline (2025–2026)", fontsize=13, fontweight="bold", pad=12)
  ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f"${x:.0f}B"))
  ax.spines["top"].set_visible(False)
  ax.spines["right"].set_visible(False)
  ax.grid(axis="y", alpha=0.3)
  plt.tight_layout()
  path = CHARTS_DIR / "valuation_timeline.png"
  fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
  plt.close(fig)
  return path


def chart_subscriber_growth():
  years = ["2023", "2024", "Dec 2025", "Feb 2026"]
  subscribers = [2.3, 4.6, 9.0, 10.0]
  fig, ax = plt.subplots(figsize=(8, 4.5))
  ax.bar(years, subscribers, color="#06B6D4", edgecolor="white", linewidth=1.2, width=0.55)
  for i, v in enumerate(subscribers):
    ax.text(i, v + 0.2, f"{v}M", ha="center", fontweight="bold", fontsize=10)
  ax.set_ylabel("Starlink Subscribers (Millions)", fontsize=11)
  ax.set_title("Starlink Subscriber Growth", fontsize=13, fontweight="bold", pad=12)
  ax.set_ylim(0, 12)
  ax.spines["top"].set_visible(False)
  ax.spines["right"].set_visible(False)
  ax.grid(axis="y", alpha=0.3)
  plt.tight_layout()
  path = CHARTS_DIR / "subscriber_growth.png"
  fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
  plt.close(fig)
  return path


def add_figure(doc, image_path, caption, width=5.8):
  doc.add_picture(str(image_path), width=Inches(width))
  last = doc.paragraphs[-1]
  last.alignment = WD_ALIGN_PARAGRAPH.CENTER
  cap = doc.add_paragraph(caption)
  cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
  cap.runs[0].italic = True
  cap.runs[0].font.size = Pt(9)
  cap.runs[0].font.color.rgb = RGBColor(80, 80, 80)
  doc.add_paragraph()


def build_report():
  CHARTS_DIR.mkdir(exist_ok=True)

  stock_chart = chart_stock_price()
  revenue_chart = chart_revenue_breakdown()
  valuation_chart = chart_valuation_timeline()
  subscriber_chart = chart_subscriber_growth()

  doc = Document()

  title = doc.add_heading("SpaceX Stock Information Report", 0)
  title.alignment = WD_ALIGN_PARAGRAPH.CENTER

  sub = doc.add_paragraph("Ticker: NASDAQ: SPCX  |  Report Date: June 15, 2026")
  sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
  sub.runs[0].font.size = Pt(11)
  sub.runs[0].font.color.rgb = RGBColor(100, 100, 100)

  doc.add_paragraph()

  doc.add_heading("Executive Summary", level=1)
  doc.add_paragraph(
    "Space Exploration Technologies Corp. (SpaceX) completed the largest initial public offering "
    "in history on June 12, 2026, listing on the Nasdaq under ticker symbol SPCX. The company "
    "priced its IPO at $135 per share, raising $75 billion and achieving an initial valuation of "
    "approximately $1.75 trillion. On its first trading day, shares surged 19.2% to close at "
    "$160.95, pushing market capitalization above $2 trillion. As of June 15, 2026, SPCX trades "
    "near $171.91 per share."
  )

  doc.add_heading("1. IPO Overview", level=1)
  add_table(doc,
    ["Metric", "Value"],
    [
      ["Ticker Symbol", "SPCX (NASDAQ)"],
      ["IPO Date", "June 12, 2026"],
      ["IPO Price", "$135.00 per share"],
      ["Shares Offered", "555.5 million"],
      ["Capital Raised", "$75 billion"],
      ["IPO Valuation", "$1.75 trillion"],
      ["First-Day Open", "$150.00"],
      ["First-Day Close", "$160.95 (+19.2%)"],
      ["Current Price (Jun 15)", "$171.91"],
      ["Market Capitalization", "~$2.1 trillion"],
      ["Lead Underwriters", "Morgan Stanley, Goldman Sachs, JPMorgan"],
      ["Retail Allocation", "~20–22% of shares"],
    ]
  )
  doc.add_paragraph()

  doc.add_heading("2. Stock Price Performance", level=1)
  doc.add_paragraph(
    "Figure 1 illustrates SPCX price movement from the final private-market estimate on June 11 "
    "through the IPO debut and current trading levels. The stock opened above its IPO price and "
    "has continued to attract strong institutional demand in pre-market and regular sessions."
  )
  add_figure(doc, stock_chart, "Figure 1: SPCX share price during IPO week (June 2026)")

  doc.add_heading("3. Financial Performance (2025)", level=1)
  add_table(doc,
    ["Financial Metric", "2024", "2025", "YoY Change"],
    [
      ["Total Revenue", "$14.1B", "$18.7B", "+33%"],
      ["Starlink Revenue", "$7.7B", "$11.4B", "+48%"],
      ["Adjusted EBITDA", "$4.9B", "$6.6B", "+35%"],
      ["GAAP Net Income", "—", "-$4.9B", "Loss"],
      ["Starlink Operating Profit", "$2.4B", "$4.4B", "+83%"],
      ["Starlink Subscribers", "4.6M", "9.0M", "+96%"],
      ["Operational Satellites", "~7,000", "~9,000+", "+29%"],
    ]
  )
  doc.add_paragraph()

  doc.add_paragraph(
    "Starlink is SpaceX's primary revenue and profit driver, accounting for 61% of total revenue "
    "in 2025. The launch services division and emerging AI/compute initiatives (including the "
    "February 2026 xAI merger) comprise the remainder. Figure 2 shows the revenue breakdown."
  )
  add_figure(doc, revenue_chart, "Figure 2: SpaceX 2025 revenue by business segment ($ billions)")

  doc.add_heading("4. Valuation History", level=1)
  doc.add_paragraph(
    "SpaceX's valuation has grown rapidly through secondary share sales and the xAI merger "
    "before reaching its historic IPO valuation. Figure 3 tracks key valuation milestones."
  )
  add_table(doc,
    ["Date / Event", "Share Price", "Valuation"],
    [
      ["July 2025 (Secondary Sale)", "$212", "$400 billion"],
      ["December 2025 (Tender Offer)", "$421", "$800 billion"],
      ["February 2026 (xAI Merger)", "—", "$1.25 trillion"],
      ["June 12, 2026 (IPO Pricing)", "$135", "$1.75 trillion"],
      ["June 15, 2026 (Market)", "$171.91", "~$2.1 trillion"],
    ]
  )
  doc.add_paragraph()
  add_figure(doc, valuation_chart, "Figure 3: SpaceX valuation milestones from 2025 to June 2026")

  doc.add_heading("5. Starlink Growth", level=1)
  doc.add_paragraph(
    "Starlink reached approximately 10 million subscribers by February 2026, doubling its "
    "subscriber base year-over-year. Average revenue per user (ARPU) declined to roughly $81/month "
    "as lower-cost consumer plans expanded the addressable market. Figure 4 shows subscriber growth."
  )
  add_figure(doc, subscriber_chart, "Figure 4: Starlink global subscriber growth (millions)")

  doc.add_heading("6. Key Investment Considerations", level=1)
  add_table(doc,
    ["Factor", "Detail", "Implication"],
    [
      ["Profitability", "GAAP net loss of $4.9B in 2025", "Heavy capex & AI spending weigh on earnings"],
      ["Starlink", "$4.4B operating profit", "Core profit engine supports valuation"],
      ["Governance", "Class B shares (10 votes each)", "Musk retains voting control post-IPO"],
      ["Index Inclusion", "Nasdaq-100 eligible after 15 days", "Passive fund buying expected"],
      ["AI Exposure", "xAI merger + Anthropic GPU deal", "Growth catalyst but adds risk"],
      ["Valuation", "~9x revenue at IPO", "Premium vs. traditional aerospace peers"],
    ]
  )
  doc.add_paragraph()

  doc.add_heading("7. How to Invest", level=1)
  doc.add_paragraph(
    "SPCX is now publicly traded and can be purchased through any standard brokerage account "
    "(Fidelity, Schwab, Robinhood, etc.). Investors may also gain indirect exposure through ETFs "
    "such as ERShares Private-Public Crossover ETF (XOVR), Tema Space Innovators ETF (NASA), and "
    "KraneShares AI & Technology ETF (KOID), which held SpaceX positions prior to the IPO."
  )

  doc.add_heading("8. Sources", level=1)
  sources = [
    "Nasdaq Private Market — SpaceX company profile and IPO data",
    "The Motley Fool — SpaceX IPO debut analysis (June 12, 2026)",
    "MarketBeat — SPCX stock debut and sector rotation analysis",
    "Sacra — SpaceX revenue, valuation & funding estimates",
    "Morningstar — 6 Charts on SpaceX S-1 Financials",
    "BitMEX Research — SpaceX IPO Guide and S-1 breakdown",
    "SpaceX Form S-1/A filing (SEC, June 2026)",
  ]
  for src in sources:
    doc.add_paragraph(src, style="List Bullet")

  note = doc.add_paragraph()
  note.add_run("Disclaimer: ").bold = True
  note.add_run(
    "This report is for informational purposes only and does not constitute investment advice. "
    "Stock prices and financial figures are based on publicly available sources as of June 15, 2026."
  )

  doc.save(OUTPUT)
  print(f"Report saved: {OUTPUT}")
  print(f"Size: {OUTPUT.stat().st_size:,} bytes")


if __name__ == "__main__":
  build_report()
